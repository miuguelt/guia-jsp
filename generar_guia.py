"""
============================================
GENERADOR DE GUÍA DE APRENDIZAJE - SENA ADSO
Script que genera el documento Word (.docx)
de la guía de aprendizaje sobre JSP + MVC
con JDK 21, Jakarta EE y PostgreSQL.
============================================
"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
RECURSOS_DIR = os.path.join(BASE_DIR, "recursos")
CODIGO_DIR = os.path.join(RECURSOS_DIR, "codigo-ejemplo")
FASTAPI_DIR = os.path.join(RECURSOS_DIR, "fastapi-inventario")
REACT_DIR = os.path.join(RECURSOS_DIR, "frontend-inventario")


def leer_archivo(ruta_relativa):
    """Lee un archivo de recursos y retorna su contenido como texto."""
    ruta = os.path.join(CODIGO_DIR, ruta_relativa)
    with open(ruta, "r", encoding="utf-8") as f:
        return f.read()


def leer_fastapi(ruta_relativa):
    """Lee un archivo del proyecto FastAPI."""
    ruta = os.path.join(FASTAPI_DIR, ruta_relativa)
    with open(ruta, "r", encoding="utf-8") as f:
        return f.read()


def leer_react(ruta_relativa):
    """Lee un archivo del proyecto React."""
    ruta = os.path.join(REACT_DIR, ruta_relativa)
    with open(ruta, "r", encoding="utf-8") as f:
        return f.read()


def leer_sql():
    ruta = os.path.join(RECURSOS_DIR, "sql", "inventario_db.sql")
    with open(ruta, "r", encoding="utf-8") as f:
        return f.read()


def leer_dockerfile():
    ruta = os.path.join(RECURSOS_DIR, "docker", "Dockerfile")
    with open(ruta, "r", encoding="utf-8") as f:
        return f.read()


def configurar_estilos(doc):
    """Configura los estilos del documento."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    for nivel in range(1, 4):
        style_name = f"Heading {nivel}"
        if style_name in doc.styles:
            h = doc.styles[style_name]
            h.font.color.rgb = RGBColor(0x4A, 0x2C, 0x8A)
            h.font.name = "Calibri"

    try:
        code_style = doc.styles.add_style("Codigo", WD_STYLE_TYPE.PARAGRAPH)
    except ValueError:
        code_style = doc.styles["Codigo"]
    code_style.font.name = "Consolas"
    code_style.font.size = Pt(9)
    code_style.paragraph_format.space_before = Pt(0)
    code_style.paragraph_format.space_after = Pt(0)
    code_style.paragraph_format.line_spacing = 1.0


def agregar_portada(doc):
    """Agrega la portada institucional."""
    for _ in range(4):
        doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SERVICIO NACIONAL DE APRENDIZAJE - SENA")
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x4A, 0x2C, 0x8A)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Tecnólogo en Análisis y Desarrollo de Software (ADSO)")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("GUÍA DE APRENDIZAJE")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x4A, 0x2C, 0x8A)

    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Desarrollo de Aplicaciones Web con JSP\nPatrón MVC con Jakarta EE 10 y JDK 21")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    for _ in range(4):
        doc.add_paragraph("")

    tabla_info = doc.add_table(rows=5, cols=2)
    tabla_info.alignment = WD_TABLE_ALIGNMENT.CENTER
    datos = [
        ("Versión", "1.0"),
        ("Fecha", "Junio 2026"),
        ("Autor", "Instructor ADSO - SENA"),
        ("Duración estimada", "40 horas"),
        ("Modalidad", "Presencial / Virtual"),
    ]
    for i, (clave, valor) in enumerate(datos):
        celda_clave = tabla_info.rows[i].cells[0]
        celda_valor = tabla_info.rows[i].cells[1]
        run_k = celda_clave.paragraphs[0].add_run(clave)
        run_k.font.bold = True
        run_k.font.size = Pt(11)
        run_v = celda_valor.paragraphs[0].add_run(valor)
        run_v.font.size = Pt(11)

    doc.add_page_break()


def agregar_tabla_identificacion(doc):
    """Sección 1: Identificación de la Guía."""
    doc.add_heading("1. IDENTIFICACIÓN DE LA GUÍA DE APRENDIZAJE", level=1)

    tabla = doc.add_table(rows=7, cols=2)
    tabla.style = "Table Grid"
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER

    datos = [
        ("Denominación del Programa", "Análisis y Desarrollo de Software (ADSO)"),
        ("Nivel de Formación", "Técnico / Tecnólogo"),
        ("Competencia", "Desarrollar la solución de software de acuerdo con el diseño y metodologías de desarrollo establecidas"),
        ("Resultado de Aprendizaje", "Construir el aplicativo web con conexión a base de datos usando el patrón Modelo-Vista-Controlador (MVC)"),
        ("Duración", "40 horas (20 teoría + 20 práctica)"),
        ("Tecnologías", "JDK 21, Jakarta EE 10, JSP, JSTL, PostgreSQL, Maven, Tomcat 10.1, Docker"),
        ("Metodología", "Aprendizaje basado en proyectos - Patrón MVC"),
    ]
    for i, (clave, valor) in enumerate(datos):
        celda_clave = tabla.rows[i].cells[0]
        celda_valor = tabla.rows[i].cells[1]
        run_k = celda_clave.paragraphs[0].add_run(clave)
        run_k.font.bold = True
        run_k.font.size = Pt(10)
        run_v = celda_valor.paragraphs[0].add_run(valor)
        run_v.font.size = Pt(10)
        celda_clave.width = Cm(5)
        celda_valor.width = Cm(12)

    doc.add_paragraph("")
    doc.add_page_break()


def agregar_presentacion(doc):
    """Sección 2: Presentación."""
    doc.add_heading("2. PRESENTACIÓN", level=1)

    doc.add_paragraph(
        "¡Bienvenido(a), Aprendiz! En la era digital actual, el desarrollo de aplicaciones web "
        "robustas, escalables y mantenibles es una de las competencias más demandadas en el "
        "mercado laboral. Esta guía de aprendizaje te llevará paso a paso en la construcción "
        "de un Sistema de Inventario completo utilizando las tecnologías más relevantes del "
        "ecosistema Java moderno."
    )

    doc.add_heading("¿Por qué el Patrón MVC?", level=2)
    doc.add_paragraph(
        "El patrón Modelo-Vista-Controlador (MVC) es el estándar de facto en el desarrollo "
        "web profesional. Separar tu aplicación en estas tres capas te permite:"
    )
    items_mvc = [
        "Mantenibilidad: Cambiar la interfaz no afecta la lógica de negocio.",
        "Testabilidad: Cada componente puede probarse de forma independiente.",
        "Trabajo en equipo: Diferentes desarrolladores pueden trabajar en capas distintas.",
        "Reutilización: Los modelos y la lógica de acceso a datos se reutilizan fácilmente.",
    ]
    for item in items_mvc:
        p = doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("¿Por qué JDK 21 y Jakarta EE?", level=2)
    doc.add_paragraph(
        "JDK 21 es la versión LTS (Long Term Support) más reciente de Java, que incluye "
        "características revolucionarias como los Records (clases inmutables con sintaxis "
        "concisa), Pattern Matching, Virtual Threads y más. Jakarta EE (sucesor de Java EE) "
        "es el estándar para aplicaciones empresariales en Java, usando el namespace "
        "jakarta.servlet.* en lugar del antiguo javax.servlet.*."
    )

    doc.add_heading("¿Por qué Docker y Coolify?", level=2)
    doc.add_paragraph(
        "El despliegue moderno de aplicaciones requiere conocimientos de contenedores. "
        "Docker permite empaquetar tu aplicación con todas sus dependencias en una imagen "
        "que funciona igual en cualquier entorno. Coolify es una plataforma de despliegue "
        "autohospedada que facilita la puesta en producción de tus aplicaciones."
    )

    doc.add_paragraph(
        "A lo largo de esta guía, construirás un sistema funcional desde cero, aplicando "
        "buenas prácticas de programación orientada a objetos (POO), seguridad y arquitectura "
        "de software. ¡Manos a la obra!"
    )
    doc.add_page_break()


def agregar_reflexion(doc):
    """Sección 3.1: Actividad de Reflexión Inicial."""
    doc.add_heading("3. FORMULACIÓN DE LAS ACTIVIDADES DE APRENDIZAJE", level=1)
    doc.add_heading("3.1 Actividad de Reflexión Inicial", level=2)

    doc.add_paragraph(
        "Antes de comenzar a programar, reflexionemos sobre un problema común en el "
        "desarrollo web. Lee el siguiente caso de estudio:"
    )

    doc.add_heading("Caso de Estudio: El Código Espagueti", level=3)

    p = doc.add_paragraph()
    run = p.add_run(
        "Imagina que eres el nuevo desarrollador en una empresa. Tu jefe te pide modificar "
        "una página JSP que muestra una lista de productos. Al abrir el archivo, encuentras "
        "algo así:"
    )

    codigo_espagueti = (
        '<%@ page import="java.sql.*" %>\n'
        "<%\n"
        '    String url = "jdbc:postgresql://localhost:5432/mi_db";\n'
        '    String user = "postgres";\n'
        '    String pass = "mi_password_secreta";\n'
        "    Connection conn = DriverManager.getConnection(url, user, pass);\n"
        '    String sql = "SELECT * FROM productos WHERE precio > " + request.getParameter("precio");\n'
        "    Statement stmt = conn.createStatement();\n"
        "    ResultSet rs = stmt.executeQuery(sql);\n"
        "%>\n"
        "<html>\n"
        "<body>\n"
        "    <table>\n"
        "    <% while(rs.next()) { %>\n"
        "        <tr>\n"
        "            <td><%= rs.getString(\"nombre\") %></td>\n"
        "            <td><%= rs.getDouble(\"precio\") %></td>\n"
        "        </tr>\n"
        "    <% } %>\n"
        "    </table>\n"
        "    <%\n"
        '    if (request.getParameter("accion") != null && request.getParameter("accion").equals("eliminar")) {\n'
        '        String id = request.getParameter("id");\n'
        '        conn.createStatement().executeUpdate("DELETE FROM productos WHERE id = " + id);\n'
        "    }\n"
        "    %>\n"
        "</body>\n"
        "</html>"
    )
    agregar_bloque_codigo(doc, codigo_espagueti)

    doc.add_heading("Preguntas de Reflexión:", level=3)
    preguntas = [
        "¿Puedes identificar al menos 3 problemas de seguridad en este código?",
        "¿Qué pasaría si necesitas cambiar la base de datos de PostgreSQL a MySQL? ¿Cuántos archivos tendrías que modificar?",
        "¿Cómo harías para probar la lógica de negocio sin tener un servidor web funcionando?",
        "Si dos desarrolladores necesitan trabajar en este archivo al mismo tiempo, ¿qué problemas podrían surgir?",
        "¿Por qué el patrón MVC soluciona todos estos problemas?",
    ]
    for i, pregunta in enumerate(preguntas, 1):
        doc.add_paragraph(f"{i}. {pregunta}")

    doc.add_paragraph(
        "Discute tus respuestas con tu equipo de trabajo y con tu instructor. "
        "Al finalizar esta guía, tendrás las herramientas para resolver cada uno de estos problemas."
    )
    doc.add_page_break()


def agregar_apropiacion(doc):
    """Sección 3.2: Apropiación del Conocimiento."""
    doc.add_heading("3.2 Actividad de Apropiación del Conocimiento", level=2)
    doc.add_paragraph(
        "En esta sección construiremos paso a paso nuestro Sistema de Inventario. "
        "Sigue cada paso cuidadosamente y asegúrate de comprender el propósito de cada componente."
    )

    # --- PASO 1: CONFIGURACIÓN DEL ENTORNO ---
    doc.add_heading("Paso 1: Configuración del Entorno de Desarrollo", level=3)

    doc.add_heading("Herramientas necesarias:", level=4)
    tabla = doc.add_table(rows=7, cols=3)
    tabla.style = "Table Grid"
    headers = ["Herramienta", "Versión", "Propósito"]
    for i, h in enumerate(headers):
        run = tabla.rows[0].cells[i].paragraphs[0].add_run(h)
        run.font.bold = True

    datos_tabla = [
        ("JDK", "21 (LTS)", "Compilación y ejecución de Java"),
        ("Apache Maven", "3.9+", "Gestión de dependencias y build"),
        ("Apache Tomcat", "10.1.x", "Servidor de aplicaciones Jakarta EE"),
        ("PostgreSQL", "12+", "Base de datos relacional"),
        ("IDE", "IntelliJ / Eclipse / VS Code", "Entorno de desarrollo"),
        ("Docker", "24+", "Contenedores para despliegue"),
    ]
    for i, (herr, ver, prop) in enumerate(datos_tabla, 1):
        tabla.rows[i].cells[0].paragraphs[0].add_run(herr)
        tabla.rows[i].cells[1].paragraphs[0].add_run(ver)
        tabla.rows[i].cells[2].paragraphs[0].add_run(prop)

    doc.add_paragraph("")
    doc.add_heading("Verificar instalación de JDK 21:", level=4)
    agregar_bloque_codigo(doc, "java --version\n# Debe mostrar: java 21.x.x")

    doc.add_heading("Iniciar servidor Tomcat 10.1:", level=4)
    agregar_bloque_codigo(
        doc,
        "# En Windows:\n"
        "cd C:\\apache-tomcat-10.1.x\\bin\n"
        "startup.bat\n\n"
        "# En Linux/Mac:\n"
        "cd /opt/apache-tomcat-10.1.x/bin\n"
        "./startup.sh\n\n"
        "# Verificar en: http://localhost:8080",
    )

    # --- PASO 2: BASE DE DATOS ---
    doc.add_heading("Paso 2: Creación de la Base de Datos", level=3)
    doc.add_paragraph(
        "Ejecuta el siguiente script SQL en tu servidor PostgreSQL para crear la base de datos "
        "y las tablas necesarias. El script incluye datos de prueba para que puedas comenzar "
        "a trabajar inmediatamente."
    )

    agregar_bloque_codigo(doc, "-- Crear la base de datos\nCREATE DATABASE inventario_db;")
    doc.add_paragraph("")
    doc.add_paragraph(
        "A continuación, el script completo para crear las tablas (roles, usuarios, productos):"
    )
    agregar_bloque_codigo(doc, leer_sql())

    # --- PASO 3: ESTRUCTURA DEL PROYECTO ---
    doc.add_heading("Paso 3: Estructura del Proyecto Maven", level=3)
    doc.add_paragraph(
        "Creamos un proyecto Maven con la siguiente estructura de directorios. "
        "Esta organización sigue el estándar Maven para aplicaciones web (WAR):"
    )

    estructura = (
        "inventario-mvc/\n"
        "├── pom.xml\n"
        "├── src/\n"
        "│   ├── main/\n"
        "│   │   ├── java/com/sena/inventario/\n"
        "│   │   │   ├── modelo/        → Entidades (Records)\n"
        "│   │   │   ├── dao/           → Acceso a datos (JDBC)\n"
        "│   │   │   ├── controlador/   → Servlets\n"
        "│   │   │   └── util/          → Utilidades\n"
        "│   │   ├── resources/\n"
        "│   │   │   └── db.properties  → Configuración BD\n"
        "│   │   └── webapp/\n"
        "│   │       ├── WEB-INF/web.xml\n"
        "│   │       ├── productos/     → Vistas JSP de productos\n"
        "│   │       ├── usuarios/      → Vistas JSP de usuarios\n"
        "│   │       ├── css/           → Hojas de estilo\n"
        "│   │       ├── index.jsp\n"
        "│   │       └── login.jsp\n"
        "│   └── test/\n"
        "└── Dockerfile"
    )
    agregar_bloque_codigo(doc, estructura)

    doc.add_heading("Archivo pom.xml:", level=4)
    doc.add_paragraph(
        "Este archivo define las dependencias del proyecto. Observa que usamos "
        "jakarta.servlet-api (no javax.servlet) y el driver de PostgreSQL:"
    )
    agregar_bloque_codigo(doc, leer_archivo("pom.xml"))

    # --- PASO 4: MODELO ---
    doc.add_heading("Paso 4: El Modelo (Model) - Entidades con Records de JDK 21", level=3)
    doc.add_paragraph(
        "Los records son una característica de JDK 14+ (estable desde JDK 16) que nos permite "
        "crear clases inmutables con una sintaxis muy concisa. Automáticamente generan "
        "constructor, getters, equals(), hashCode() y toString()."
    )

    doc.add_heading("Rol.java - Record para la entidad Rol:", level=4)
    agregar_bloque_codigo(doc, leer_archivo(os.path.join("src", "main", "java", "com", "sena", "inventario", "modelo", "Rol.java")))

    doc.add_heading("Usuario.java - Record para la entidad Usuario:", level=4)
    agregar_bloque_codigo(doc, leer_archivo(os.path.join("src", "main", "java", "com", "sena", "inventario", "modelo", "Usuario.java")))

    doc.add_heading("Producto.java - Record para la entidad Producto:", level=4)
    agregar_bloque_codigo(doc, leer_archivo(os.path.join("src", "main", "java", "com", "sena", "inventario", "modelo", "Producto.java")))

    # --- PASO 5: DAO ---
    doc.add_heading("Paso 5: Capa de Acceso a Datos (DAO) con JDBC", level=3)
    doc.add_paragraph(
        "El patrón DAO (Data Access Object) separa la lógica de acceso a datos de la lógica "
        "de negocio. Cada entidad tiene su propia clase DAO que encapsula todas las operaciones "
        "CRUD (Create, Read, Update, Delete)."
    )

    doc.add_heading("ConexionDB.java - Gestión de conexiones:", level=4)
    agregar_bloque_codigo(doc, leer_archivo(os.path.join("src", "main", "java", "com", "sena", "inventario", "dao", "ConexionDB.java")))

    doc.add_heading("ProductoDAO.java - Operaciones CRUD completas:", level=4)
    doc.add_paragraph(
        "Esta clase implementa las 4 operaciones CRUD para la entidad Producto. "
        "Observa el uso de PreparedStatement para prevenir inyecciones SQL y "
        "try-with-resources para garantizar el cierre de recursos:"
    )
    agregar_bloque_codigo(doc, leer_archivo(os.path.join("src", "main", "java", "com", "sena", "inventario", "dao", "ProductoDAO.java")))

    doc.add_heading("UsuarioDAO.java - Operaciones CRUD:", level=4)
    agregar_bloque_codigo(doc, leer_archivo(os.path.join("src", "main", "java", "com", "sena", "inventario", "dao", "UsuarioDAO.java")))

    doc.add_heading("RolDAO.java - Operaciones CRUD:", level=4)
    agregar_bloque_codigo(doc, leer_archivo(os.path.join("src", "main", "java", "com", "sena", "inventario", "dao", "RolDAO.java")))

    # --- PASO 6: CONTROLADOR ---
    doc.add_heading("Paso 6: El Controlador (Controller) - Servlets", level=3)
    doc.add_paragraph(
        "Los Servlets son las clases Java que manejan las peticiones HTTP. Actúan como "
        "coordinadores: reciben la petición del usuario, invocan la lógica del Modelo y "
        "seleccionan la Vista apropiada para mostrar la respuesta."
    )

    doc.add_heading("ProductoServlet.java - Controlador de Productos:", level=4)
    doc.add_paragraph(
        "Este servlet maneja todas las operaciones CRUD de productos. "
        "Usa la anotación @WebServlet para mapear URLs y el patrón switch con "
        "flecha (->) de JDK 14+ para enrutar las acciones:"
    )
    agregar_bloque_codigo(doc, leer_archivo(os.path.join("src", "main", "java", "com", "sena", "inventario", "controlador", "ProductoServlet.java")))

    doc.add_heading("UsuarioServlet.java - Controlador de Usuarios:", level=4)
    agregar_bloque_codigo(doc, leer_archivo(os.path.join("src", "main", "java", "com", "sena", "inventario", "controlador", "UsuarioServlet.java")))

    doc.add_heading("LoginServlet.java - Controlador de Autenticación:", level=4)
    agregar_bloque_codigo(doc, leer_archivo(os.path.join("src", "main", "java", "com", "sena", "inventario", "controlador", "LoginServlet.java")))

    doc.add_heading("PasswordUtil.java - Utilidad de encriptación BCrypt:", level=4)
    agregar_bloque_codigo(doc, leer_archivo(os.path.join("src", "main", "java", "com", "sena", "inventario", "util", "PasswordUtil.java")))

    # --- PASO 7: VISTA ---
    doc.add_heading("Paso 7: La Vista (View) - Páginas JSP con JSTL", level=3)
    doc.add_paragraph(
        "Las vistas JSP usan Jakarta JSTL (JSP Standard Tag Library) para evitar "
        "código Java incrustado. Las etiquetas <c:forEach>, <c:if>, <c:out> y "
        "Expression Language (${...}) nos permiten mostrar datos de forma limpia."
    )

    doc.add_paragraph(
        "IMPORTANTE: Nota que usamos la URI jakarta.tags.core en lugar de la antigua "
        "http://java.sun.com/jsp/jstl/core. Esto es parte de la migración a Jakarta EE."
    )

    doc.add_heading("lista.jsp - Vista de lista de productos con JSTL:", level=4)
    agregar_bloque_codigo(doc, leer_archivo(os.path.join("src", "main", "webapp", "productos", "lista.jsp")))

    doc.add_heading("formulario.jsp - Vista de formulario para crear/editar:", level=4)
    agregar_bloque_codigo(doc, leer_archivo(os.path.join("src", "main", "webapp", "productos", "formulario.jsp")))

    # --- PASO 8: COMPILACIÓN ---
    doc.add_heading("Paso 8: Compilación y Ejecución", level=3)
    doc.add_paragraph("Para compilar y empaquetar la aplicación, ejecuta:")
    agregar_bloque_codigo(doc, "# Compilar y generar el archivo WAR\nmvn clean package")
    doc.add_paragraph(
        "El archivo WAR se generará en target/inventario-mvc-1.0.war. "
        "Para desplegarlo en Tomcat, cópialo a la carpeta webapps/ del servidor."
    )
    agregar_bloque_codigo(
        doc,
        "# Copiar WAR a Tomcat\ncopy target\\inventario-mvc-1.0.war C:\\apache-tomcat-10.1.x\\webapps\\\n\n"
        "# Reiniciar Tomcat\nC:\\apache-tomcat-10.1.x\\bin\\shutdown.bat\nC:\\apache-tomcat-10.1.x\\bin\\startup.bat\n\n"
        "# Acceder a la aplicación\n# http://localhost:8080/inventario-mvc-1.0/",
    )

    doc.add_page_break()

    # --- 3.3 TRANSFERENCIA ---
    doc.add_heading("3.3 Actividad de Transferencia del Conocimiento", level=2)
    doc.add_paragraph(
        "Ahora que has comprendido y replicado el código base, es momento de "
        "demostrar tu dominio aplicando lo aprendido en nuevos escenarios. "
        "Resuelve los siguientes ejercicios retadores:"
    )

    doc.add_heading("Ejercicio 1: Buscador de Productos", level=3)
    doc.add_paragraph(
        "Modifica el sistema para agregar un buscador de productos por nombre. "
        "El aprendiz debe:"
    )
    items_ej1 = [
        "Agregar un campo de búsqueda en la vista lista.jsp con un formulario GET.",
        "Modificar el ProductoDAO para incluir un método buscarPorNombre(String nombre) que use LIKE en SQL.",
        "Agregar una nueva acción 'buscar' en el ProductoServlet que invoque el nuevo método DAO.",
        "Mostrar los resultados filtrados en la misma vista lista.jsp, indicando cuántos resultados se encontraron.",
        "Bonus: Agregar búsqueda por rango de precios (precio mínimo y máximo).",
    ]
    for item in items_ej1:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Ejercicio 2: Tabla Categorías con Relación a Productos", level=3)
    doc.add_paragraph(
        "Crea una nueva entidad Categoria y enlázala con Productos mediante una relación 1:N. "
        "El aprendiz debe:"
    )
    items_ej2 = [
        "Crear la tabla categorias en PostgreSQL (id, nombre, descripcion, estado).",
        "Crear el record Categoria.java en el modelo.",
        "Crear la clase CategoriaDAO.java con operaciones CRUD completas.",
        "Modificar la tabla productos para agregar categoria_id como FOREIGN KEY.",
        "Actualizar el record Producto para incluir categoryId en lugar del String categoria.",
        "Crear CategoriaServlet.java con CRUD completo y sus vistas JSP.",
        "Modificar el formulario de productos para que use un <select> con las categorías de la BD.",
        "Bonus: Implementar eliminación en cascada o restricción si hay productos asociados.",
    ]
    for item in items_ej2:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Ejercicio 3: Sistema de Auditoría con Filtros", level=3)
    doc.add_paragraph(
        "Implementa un sistema de auditoría que registre todas las operaciones CRUD. "
        "El aprendiz debe:"
    )
    items_ej3 = [
        "Crear la tabla auditoria en PostgreSQL (id, tabla_afectada, operacion, usuario, fecha, datos_anteriores, datos_nuevos).",
        "Crear el record Auditoria.java y su DAO.",
        "Implementar un Jakarta Servlet Filter que registre la fecha/hora y usuario de cada petición.",
        "Modificar los métodos DAO de inserción, actualización y eliminación para que registren en auditoría.",
        "Crear una vista JSP que muestre el historial de auditoría con filtros por fecha y tipo de operación.",
        "Bonus: Agregar paginación a la lista de auditoría (mostrar 10 registros por página).",
    ]
    for item in items_ej3:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_page_break()


def agregar_fastapi(doc):
    """Sección 3.5: API REST con FastAPI."""
    doc.add_heading("3.5 API REST CON FASTAPI + SQLALCHEMY", level=1)
    doc.add_paragraph(
        "FastAPI es un framework moderno para construir APIs REST con Python 3.10+."
        " Es la alternativa directa a Jakarta EE + Servlets, pero con sintaxis Python moderna."
    )

    doc.add_heading("Primer Ejemplo: Hola Mundo FastAPI", level=2)
    doc.add_paragraph("Crea tu primer endpoint API REST en 2 minutos:")
    codigo_main = leer_fastapi("main.py")
    agregar_bloque_codigo(doc, codigo_main)

    doc.add_heading("Ejecutar y probar", level=2)
    doc.add_paragraph("Comandos para ejecutar:")
    agregar_bloque_codigo(doc, (
        "pip install fastapi uvicorn\n"
        "uvicorn main:app --reload --port 8000\n\n"
        "# Probar:\n"
        "curl http://localhost:8000\n"
        "# {\"mensaje\": \"Hola Mundo desde FastAPI!\"}\n\n"
        "# Swagger UI:\n"
        "# http://localhost:8000/docs"
    ))

    doc.add_heading("Estructura del proyecto", level=2)
    doc.add_paragraph(
        "El proyecto fastapi-inventario sigue el mismo patrón MVC que JSP:"
    )
    agregar_bloque_codigo(doc, (
        "fastapi-inventario/\n"
        "├── main.py              # Punto de entrada (como index.jsp)\n"
        "├── database.py          # Conexión SQLAlchemy (como ConexionDB.java)\n"
        "├── models/producto.py   # Modelo SQLAlchemy (como el Record)\n"
        "├── schemas/producto.py  # Esquema Pydantic (validación)\n"
        "├── routers/producto.py  # CRUD endpoints (como ProductoServlet)\n"
        "└── requirements.txt     # Dependencias (como pom.xml)"
    ))

    doc.add_paragraph(
        "Los archivos completos están disponibles en recursos/fastapi-inventario/"
    )


def agregar_react(doc):
    """Sección 3.6: Frontend con React 19."""
    doc.add_heading("3.6 FRONTEND CON REACT 19 + VITE", level=1)
    doc.add_paragraph(
        "React es una biblioteca de JavaScript para construir interfaces de usuario."
        " Es la alternativa moderna a las vistas JSP: mientras JSP renderiza HTML en el servidor,"
        " React renderiza componentes en el navegador y se comunica con FastAPI vía HTTP."
    )

    doc.add_heading("Primer Ejemplo: Componente React", level=2)
    doc.add_paragraph("Tu primer componente que consume la API de FastAPI:")
    codigo_app = leer_react("src/App.jsx")
    agregar_bloque_codigo(doc, codigo_app)

    doc.add_heading("Estructura del proyecto", level=2)
    agregar_bloque_codigo(doc, (
        "frontend-inventario/\n"
        "├── src/\n"
        "│   ├── services/api.js          # Axios (como ProductoDAO)\n"
        "│   ├── components/Navbar.jsx    # Navegación\n"
        "│   ├── components/ProductoCard.jsx  # Tarjeta (como fragmento JSP)\n"
        "│   ├── components/ProductoLista.jsx # Lista (como lista.jsp)\n"
        "│   ├── components/ProductoForm.jsx  # Formulario (como formulario.jsp)\n"
        "│   └── App.jsx                  # Raíz + Rutas\n"
        "├── package.json\n"
        "└── vite.config.js"
    ))

    doc.add_paragraph(
        "Los archivos completos están disponibles en recursos/frontend-inventario/"
    )


def agregar_docker_compose_fullstack(doc):
    """Apéndice: Docker Compose para Full Stack."""
    doc.add_heading("APÉNDICE: DOCKER COMPOSE FULL STACK", level=1)
    doc.add_paragraph(
        "Para levantar todo el stack moderno (PostgreSQL + FastAPI + React) con un solo comando:"
    )
    agregar_bloque_codigo(doc, (
        "docker compose -f docker-compose.fullstack.yml up -d\n\n"
        "# Servicios:\n"
        "# - PostgreSQL 16 en :5432\n"
        "# - FastAPI en :8000 (docs en /docs)\n"
        "# - React (Nginx) en :80\n\n"
        "# Detener:\n"
        "docker compose -f docker-compose.fullstack.yml down"
    ))


def agregar_despliegue(doc):
    """Sección 4: Despliegue en Coolify."""
    doc.add_heading("4. DESPLIEGUE EN COOLIFY CON DOCKER", level=1)

    doc.add_paragraph(
        "Una vez que tu aplicación funciona correctamente en local, es momento de "
        "desplegarla en un entorno de producción. Usaremos Docker para empaquetar "
        "la aplicación y Coolify (o cualquier VPS con Docker) para alojarla."
    )

    doc.add_heading("4.1 Generación del archivo WAR", level=2)
    doc.add_paragraph("Genera el archivo WAR ejecutable con Maven:")
    agregar_bloque_codigo(doc, "mvn clean package -DskipTests")
    doc.add_paragraph(
        "Esto generará el archivo target/inventario-mvc-1.0.war que contiene "
        "toda la aplicación empaquetada y lista para desplegar."
    )

    doc.add_heading("4.2 Creación del Dockerfile", level=2)
    doc.add_paragraph(
        "El Dockerfile define cómo construir la imagen Docker de nuestra aplicación. "
        "Usamos una imagen base de Tomcat 10.1 con JDK 21 (Eclipse Temurin):"
    )
    agregar_bloque_codigo(doc, leer_dockerfile())

    doc.add_heading("4.3 Construcción de la imagen Docker", level=2)
    agregar_bloque_codigo(
        doc,
        "# Construir la imagen\n"
        "docker build -t inventario-mvc:1.0 .\n\n"
        "# Verificar que la imagen se creó\n"
        "docker images | findstr inventario-mvc",
    )

    doc.add_heading("4.4 Ejecución del contenedor", level=2)
    agregar_bloque_codigo(
        doc,
        "# Ejecutar el contenedor con variables de entorno\n"
        "docker run -d -p 8080:8080 \\\n"
        "  -e DB_HOST=host.docker.internal \\\n"
        "  -e DB_PORT=5432 \\\n"
        "  -e DB_NAME=inventario_db \\\n"
        "  -e DB_USER=postgres \\\n"
        "  -e DB_PASSWORD=tu_password_seguro \\\n"
        "  --name inventario-app \\\n"
        "  inventario-mvc:1.0\n\n"
        "# Verificar que el contenedor está corriendo\n"
        "docker ps\n\n"
        "# Ver logs en tiempo real\n"
        "docker logs -f inventario-app",
    )

    doc.add_heading("4.5 Despliegue en Coolify", level=2)
    doc.add_paragraph(
        "Coolify es una plataforma de despliegue autohospedada (alternativa open-source "
        "a Heroku/Vercel). Para desplegar tu aplicación:"
    )
    pasos_coolify = [
        "Sube tu código fuente a un repositorio de GitHub (incluye el Dockerfile).",
        "Instala Coolify en tu VPS o servidor (https://coolify.io).",
        "En el panel de Coolify, crea una nueva 'Application'.",
        "Conecta tu repositorio de GitHub y selecciona la rama main.",
        "Coolify detectará automáticamente el Dockerfile y configurará el build.",
        "Configura las variables de entorno (DB_HOST, DB_PORT, etc.) en el panel.",
        "Haz clic en 'Deploy' y espera a que la aplicación esté disponible.",
        "Accede a tu aplicación en la URL proporcionada por Coolify.",
    ]
    for i, paso in enumerate(pasos_coolify, 1):
        doc.add_paragraph(f"{i}. {paso}")

    doc.add_heading("4.6 Alternativa: Docker Compose", level=2)
    doc.add_paragraph(
        "Para facilitar el despliegue, puedes usar Docker Compose para levantar "
        "tanto la aplicación como la base de datos en un solo comando:"
    )
    docker_compose = (
        'version: "3.8"\n'
        "services:\n"
        "  db:\n"
        "    image: postgres:16-alpine\n"
        "    environment:\n"
        "      POSTGRES_DB: inventario_db\n"
        "      POSTGRES_USER: postgres\n"
        "      POSTGRES_PASSWORD: tu_password_seguro\n"
        "    ports:\n"
        '      - "5432:5432"\n'
        "    volumes:\n"
        "      - pgdata:/var/lib/postgresql/data\n"
        "      - ./recursos/sql/inventario_db.sql:/docker-entrypoint-initdb.d/init.sql\n\n"
        "  app:\n"
        "    build: .\n"
        "    ports:\n"
        '      - "8080:8080"\n'
        "    environment:\n"
        "      DB_HOST: db\n"
        "      DB_PORT: 5432\n"
        "      DB_NAME: inventario_db\n"
        "      DB_USER: postgres\n"
        "      DB_PASSWORD: tu_password_seguro\n"
        "    depends_on:\n"
        "      - db\n\n"
        "volumes:\n"
        "  pgdata:"
    )
    agregar_bloque_codigo(doc, docker_compose)
    doc.add_paragraph("")
    agregar_bloque_codigo(doc, "# Levantar todo con un solo comando\ndocker compose up -d")

    doc.add_page_break()


def agregar_evidencias(doc):
    """Sección 5: Evidencias de Aprendizaje."""
    doc.add_heading("5. EVIDENCIAS DE APRENDIZAJE", level=1)
    doc.add_paragraph(
        "Para demostrar que has alcanzado los resultados de aprendizaje esperados, "
        "debes entregar las siguientes evidencias:"
    )

    doc.add_heading("5.1 Evidencia de Conocimiento", level=2)
    doc.add_paragraph(
        "Cuestionario escrito (individual) que responda las siguientes preguntas:"
    )
    preguntas = [
        "Explique con sus propias palabras qué es el patrón MVC y cuál es la responsabilidad de cada capa.",
        "¿Cuál es la diferencia entre las importaciones javax.servlet.* y jakarta.servlet.*? ¿Por qué cambió?",
        "¿Qué ventajas ofrecen los records de JDK 21 frente a las clases tradicionales para las entidades del modelo?",
        "Explique qué es una inyección SQL y cómo PreparedStatement la previene. Incluya un ejemplo.",
        "¿Qué función cumple el archivo web.xml en una aplicación Jakarta EE 10? ¿Es obligatorio?",
        "Describa el flujo completo de una petición HTTP cuando un usuario hace clic en 'Crear Producto'.",
        "¿Qué es JSTL y por qué es mejor que usar scriptlets (<% %>) en las páginas JSP?",
        "Explique la diferencia entre forward() y sendRedirect() en un Servlet.",
        "¿Qué es un Servlet Filter y para qué se utiliza? Dé un ejemplo práctico.",
        "Describa el proceso de dockerización de una aplicación web Java. ¿Qué ventajas tiene?",
    ]
    for i, p in enumerate(preguntas, 1):
        doc.add_paragraph(f"{i}. {p}")

    doc.add_heading("5.2 Evidencia de Desempeño", level=2)
    doc.add_paragraph("Video de sustentación (máximo 10 minutos) donde el aprendiz debe:")
    items_video = [
        "Mostrar la aplicación funcionando en el navegador (listar, crear, editar, eliminar registros).",
        "Explicar la arquitectura MVC implementada, señalando cada capa en el código.",
        "Demostrar el uso de JSTL en las vistas JSP (sin código Java incrustado).",
        "Explicar cómo funciona la conexión a la base de datos con JDBC.",
        "Mostrar la ejecución del script SQL y los datos de prueba.",
        "Explicar al menos 2 ejercicios de transferencia resueltos.",
        "Demostrar el despliegue con Docker (construcción de imagen y ejecución del contenedor).",
    ]
    for item in items_video:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("5.3 Evidencia de Producto", level=2)
    doc.add_paragraph("Repositorio en GitHub que contenga:")
    items_repo = [
        "Código fuente completo del proyecto Maven (estructura estándar).",
        "Archivo pom.xml con todas las dependencias configuradas.",
        "Script SQL de la base de datos (recursos/sql/inventario_db.sql).",
        "Dockerfile para la construcción de la imagen.",
        "Archivo docker-compose.yml (opcional pero recomendado).",
        "README.md con instrucciones de instalación y ejecución.",
        "Archivo WAR compilado (target/inventario-mvc-1.0.war) o instrucciones para generarlo.",
        "Capturas de pantalla de la aplicación funcionando.",
        "Evidencia de los ejercicios de transferencia resueltos.",
    ]
    for item in items_repo:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Rúbrica de Evaluación", level=3)
    tabla = doc.add_table(rows=7, cols=4)
    tabla.style = "Table Grid"
    headers = ["Criterio", "Excelente (100%)", "Aceptable (70%)", "Insuficiente (40%)"]
    for i, h in enumerate(headers):
        run = tabla.rows[0].cells[i].paragraphs[0].add_run(h)
        run.font.bold = True
        run.font.size = Pt(9)

    rubrica = [
        ("Estructura MVC", "Separación clara de capas", "Separación parcial", "Capas mezcladas"),
        ("Uso de JSTL", "Sin scriptlets, uso correcto de tags", "Algunos scriptlets menores", "Uso extensivo de scriptlets"),
        ("CRUD Completo", "3 tablas con CRUD funcional", "2 tablas con CRUD", "CRUD incompleto"),
        ("Seguridad", "PreparedStatement + BCrypt", "PreparedStatement sin BCrypt", "SQL vulnerable"),
        ("Dockerización", "Dockerfile funcional + compose", "Solo Dockerfile", "Sin Docker"),
        ("Ejercicios transferencia", "3 ejercicios completos", "2 ejercicios", "1 o ninguno"),
    ]
    for i, (crit, exc, acep, insuf) in enumerate(rubrica, 1):
        for j, val in enumerate([crit, exc, acep, insuf]):
            run = tabla.rows[i].cells[j].paragraphs[0].add_run(val)
            run.font.size = Pt(9)

    doc.add_page_break()


def agregar_flutter(doc):
    """Sección 3.7: App Móvil con Flutter."""
    doc.add_heading("3.7. APP MÓVIL CON FLUTTER + OPENAPI", level=1)

    doc.add_heading("3.7.1. ¿Qué es Flutter?", level=2)
    doc.add_paragraph(
        "Flutter es un framework de código abierto creado por Google para construir interfaces de usuario "
        "nativas y atractivas desde una sola base de código. Utiliza el lenguaje Dart y un motor de "
        "renderizado propio (Skia/Impeller) que pinta directamente los píxeles en la pantalla, logrando "
        "rendimiento nativo en iOS, Android, Web, Windows, macOS y Linux."
    )

    doc.add_heading("3.7.2. Stack Tecnológico", level=2)
    doc.add_paragraph(
        "• Flutter 3.x: Framework UI multiplataforma con Hot Reload y compilación nativa.\n"
        "• Dart 3.x: Lenguaje moderno con null safety, tipado fuerte y programación asíncrona.\n"
        "• Riverpod: Manejo de estado reactivo y tipado para aplicaciones complejas.\n"
        "• GoRouter: Navegación declarativa con soporte para rutas protegidas.\n"
        "• Dio: Cliente HTTP con interceptores para JWT, logging y manejo de errores.\n"
        "• OpenAPI Generator: Generación automática de clientes Dart desde la especificación OpenAPI."
    )

    doc.add_heading("3.7.3. Arquitectura de Capas", level=2)
    doc.add_paragraph(
        "La aplicación Flutter sigue una arquitectura limpia en capas:\n\n"
        "1. UI Layer: Widgets y pantallas que solo renderizan datos. Sin lógica de negocio.\n"
        "2. State Management: Providers de Riverpod que orquestan la comunicación entre UI y servicios.\n"
        "3. Service Layer: Casos de uso que coordinan repositorios y aplican reglas de negocio.\n"
        "4. Repository Layer: Abstracción entre datos remotos (API) y locales (caché).\n"
        "5. Data Sources: Clientes OpenAPI con Dio para llamadas HTTP, SharedPreferences para almacenamiento local."
    )

    doc.add_heading("3.7.4. Consumo de Microservicios", level=2)
    doc.add_paragraph(
        "La app Flutter se conecta a los microservicios FastAPI existentes en el proyecto. Cada servicio "
        "(Auth, Productos, Usuarios) tiene su propio dominio y base de datos. El ApiClient de Dio maneja:"
    )
    doc.add_paragraph(
        "• Inyección automática del token JWT en cada petición.\n"
        "• Renovación automática de token ante errores 401.\n"
        "• Logging de peticiones y respuestas en desarrollo.\n"
        "• Timeouts configurables para evitar bloqueos."
    )

    doc.add_heading("3.7.5. OpenAPI Generator", level=2)
    doc.add_paragraph(
        "OpenAPI Generator toma la especificación OpenAPI (openapi.json) generada por FastAPI y produce "
        "automáticamente el código Dart del cliente HTTP. Esto elimina la necesidad de escribir:"
    )
    doc.add_paragraph(
        "• Clases de modelo con fromJson/toJson.\n"
        "• Métodos HTTP para cada endpoint.\n"
        "• Serialización y deserialización JSON.\n"
        "• Manejo de errores básico."
    )
    doc.add_paragraph(
        "Comando de generación:\n"
        "npx @openapitools/openapi-generator-cli generate -i openapi.json -g dart -o lib/api"
    )

    doc.add_heading("3.7.6. Estructura del Proyecto Flutter", level=2)
    agregar_bloque_codigo(doc, """flutter-inventario/
+-- pubspec.yaml
+-- lib/
    +-- main.dart              # Entry point con ProviderScope
    +-- app.dart               # MaterialApp.router + GoRouter
    +-- core/
    |   +-- constants.dart     # URLs, storage keys
    |   +-- theme.dart         # Tema SENA Material 3
    |   +-- network/
    |       +-- api_client.dart # Dio + interceptors JWT
    |       +-- api_exceptions.dart
    +-- models/
    |   +-- producto.dart, usuario.dart, auth_response.dart
    +-- services/
    |   +-- auth_service.dart  # login, register, profile
    |   +-- producto_service.dart  # CRUD productos
    |   +-- usuario_service.dart   # CRUD usuarios
    +-- providers/
    |   +-- auth_provider.dart, productos_provider.dart
    +-- screens/
        +-- login, register, home
        +-- productos/ (list, form, detail)
        +-- usuarios/ (list, form)""")

    doc.add_heading("3.7.7. Despliegue con Docker", level=2)
    doc.add_paragraph(
        "La app Flutter puede compilarse para web y servirse con Nginx en un contenedor Docker. "
        "El Dockerfile multi-stage primero compila la app con Flutter SDK y luego sirve los archivos "
        "estáticos con Nginx Alpine, resultando en una imagen de solo ~20MB."
    )

def agregar_bibliografia(doc):
    """Sección 6: Bibliografía y Referencias."""
    doc.add_heading("6. BIBLIOGRAFÍA Y REFERENCIAS", level=1)

    referencias = [
        "Oracle. (2024). The Java Tutorials - Jakarta EE. https://jakarta.ee/",
        "Apache Software Foundation. (2024). Apache Tomcat 10.1 Documentation. https://tomcat.apache.org/tomcat-10.1-doc/",
        "Eclipse Foundation. (2024). Jakarta Servlet Specification 6.0. https://jakarta.ee/specifications/servlet/6.0/",
        "PostgreSQL Global Development Group. (2024). PostgreSQL 16 Documentation. https://www.postgresql.org/docs/16/",
        "Apache Maven. (2024). Maven - Introduction to the POM. https://maven.apache.org/guides/introduction/",
        "Docker Inc. (2024). Dockerfile reference. https://docs.docker.com/engine/reference/builder/",
        "Oracle. (2024). JDK 21 Release Notes. https://jdk.java.net/21/",
        "Coolify. (2024). Documentation. https://coolify.io/docs/",
        "Core JavaServer Faces (4th Edition) - David Geary, Cay Horstmann.",
        "Patterns of Enterprise Application Architecture - Martin Fowler.",
        "Google. (2024). Flutter Documentation. https://flutter.dev/docs",
        "Google. (2024). Dart Documentation. https://dart.dev/guides",
        "OpenAPI Initiative. (2024). OpenAPI Specification. https://openapis.org/",
        "Riverpod. (2024). Riverpod Documentation. https://riverpod.dev/",
    ]
    for ref in referencias:
        doc.add_paragraph(ref, style="List Bullet")


def agregar_bloque_codigo(doc, codigo):
    """Agrega un bloque de código con formato especial."""
    for linea in codigo.split("\n"):
        p = doc.add_paragraph()
        p.style = doc.styles["Codigo"]
        p.paragraph_format.left_indent = Cm(1)
        shading = p._element.get_or_add_pPr()
        shading_elem = shading.makeelement(qn("w:shd"), {
            qn("w:fill"): "F5F5F5",
            qn("w:val"): "clear",
        })
        shading.append(shading_elem)
        run = p.add_run(linea if linea else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(9)


def generar_documento():
    """Genera el documento Word completo."""
    doc = Document()
    configurar_estilos(doc)

    print("Generando portada...")
    agregar_portada(doc)

    print("Generando sección 1 - Identificación...")
    agregar_tabla_identificacion(doc)

    print("Generando sección 2 - Presentación...")
    agregar_presentacion(doc)

    print("Generando sección 3 - Actividades de Aprendizaje...")
    agregar_apropiacion(doc)

    print("Generando sección 3.5 - API REST con FastAPI...")
    agregar_fastapi(doc)

    print("Generando sección 3.6 - Frontend con React...")
    agregar_react(doc)

    print("Generando sección 3.7 - App Móvil con Flutter...")
    agregar_flutter(doc)

    print("Generando sección 4 - Despliegue...")
    agregar_despliegue(doc)

    print("Generando sección 5 - Evidencias...")
    agregar_evidencias(doc)

    print("Generando sección 6 - Bibliografía...")
    agregar_bibliografia(doc)

    print("Generando FastAPI + React Docker Compose...")
    agregar_docker_compose_fullstack(doc)

    ruta_salida = os.path.join(BASE_DIR, "Guia_Aprendizaje_JSP_MVC.docx")
    doc.save(ruta_salida)
    print(f"\nDocumento generado exitosamente: {ruta_salida}")
    return ruta_salida


if __name__ == "__main__":
    generar_documento()
